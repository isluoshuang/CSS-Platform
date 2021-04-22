# -*- encoding:utf-8 -*-
import hashlib
import json


from flask import Flask
from flask import make_response
from flask import redirect
from flask import render_template
from flask import request
from flask_cors import CORS

from flask import Flask, jsonify, request, render_template, send_from_directory, session
from time import time
from flask_pymongo import PyMongo, DESCENDING, ASCENDING



import sys
import pandas as pd
import jieba.analyse

import sys
import db

# import wechatsogou
import os
from hashlib import md5

import requests

from sklearn.decomposition import PCA as sklearnPCA
from matplotlib import pyplot as plt
import csv
from sklearn import preprocessing
from sklearn.preprocessing import MinMaxScaler
from sklearn.manifold import TSNE as sklearnTSNE
from sklearn.cluster import KMeans
from textblob import TextBlob

import re, collections

from flask import send_file, send_from_directory

from numpy import array, zeros, argmin, inf, equal, ndim
# from scipy.spatial.distance import cdist
from sklearn.metrics.pairwise import euclidean_distances

from sklearn.naive_bayes import MultinomialNB
from sklearn.pipeline import Pipeline
from sklearn.feature_extraction.text import TfidfVectorizer, HashingVectorizer, CountVectorizer
from sklearn import metrics
from sklearn.naive_bayes import BernoulliNB
import random
from scipy import stats

from sklearn import tree
from sklearn.ensemble import RandomForestClassifier

from sklearn.svm import SVC
from sklearn.svm import SVR

from sklearn.linear_model import LinearRegression
from sklearn.linear_model import LogisticRegression

from os import path
from PIL import Image
import numpy as np

from wordcloud import WordCloud, STOPWORDS, ImageColorGenerator

from pdfminer.pdfparser import PDFParser, PDFDocument
from pdfminer.pdfinterp import PDFResourceManager, PDFPageInterpreter
from pdfminer.layout import LAParams
from pdfminer.converter import PDFPageAggregator
from pdfminer.pdfinterp import PDFTextExtractionNotAllowed
from docx import Document
from scipy.interpolate import lagrange#拉格朗日函数
import warnings
warnings.filterwarnings("ignore")

from minepy import MINE
# import numpy as np
import pymongo
from pymongo import MongoClient
import hashlib
import sys
# user = "root"
# password = "ls269031126"
charset = "utf8"

app = Flask(__name__)

app.secret_key = 'zju'
app.config['MONGO_DBNAME'] = 'question'

app.config['MONGO_URI'] = 'mongodb://qs:double@10.72.100.5:8027/question'

app.url_map.strict_slashes = False

mongo = PyMongo(app)




# CORS(app)
# 跨域支持
@app.after_request
def after_request(resp):
    resp.headers['Access-Control-Allow-Origin'] = '*'
    return resp

app.after_request(after_request)


@app.route('/')
def index():
    return render_template("index.html")

global file_name
file_name = []
global loginName
loginName = ""

# 文件上传
@app.route("/upload/<name>", methods=['post', 'get'])
def upload(name):
    print(name)
    loginName = name
    f = request.files['file']
    print(f.filename)
    # file_name.append(f.filename)
    f.save(name + '/'+ f.filename)
    state = db.uploadFile(name, f.filename)
    print(state)
    print(loginName)
    # return 'http://0.0.0.0/'
    return 'https://mo.zju.edu.cn/css/'

#PCA降维
@app.route('/pca/<name>', methods=['post', 'get'])
def pca(name):
    pca_dim = request.get_json() #bytes
    data=[]
    traffic_feature=[]
    traffic_target=[]
    # fileN = file_name.pop() 
    # print(name)
    fileN = db.searchFile(name)
    # print(fileN)
    fileN = fileN[0]['filename']
    # print(fileN)
    if fileN == "dont_have_file":
        ret = {"route": "nofile"}
        return json.dumps(ret)
    csvFile = open(name + '/'+ fileN)
    csv_file = csv.reader(csvFile)
    for content in csv_file:
        content=list(map(float,content))
        if len(content)!=0:
            data.append(content)
            traffic_feature.append(content[0:-2])
            traffic_target.append(content[-1])
    csvFile.close()
    min_max_scaler = preprocessing.MinMaxScaler()
    traffic_feature = min_max_scaler.fit_transform(traffic_feature)
    # print('data=',data)
    # print('traffic_feature=',traffic_feature)
    # print('traffic_target=',traffic_target)
    dim = int(pca_dim["pca"])
    sklearn_pca = sklearnPCA(n_components = dim)
    sklearn_transf = sklearn_pca.fit_transform(traffic_feature)
    # print(sklearn_transf)   
    with open(name + '/' + 'PCA_'+fileN, 'w', newline='') as new_file:
         csv_writer = csv.writer(new_file)
         for line in sklearn_transf:
             # print(line)
             csv_writer.writerow(line)     
    ret = {"route": 'PCA_'+fileN}
    return json.dumps(ret)

#t-SNE降维
@app.route('/TSNE/<name>', methods=['post', 'get'])
def TSNE(name):
    TSNE_dim = request.get_json() #bytes
    data=[]
    traffic_feature=[]
    traffic_target=[]
    # fileN = file_name.pop() 
    # print(name)
    fileN = db.searchFile(name)
    # print(fileN)
    fileN = fileN[0]['filename']
    # print(fileN)
    if fileN == "dont_have_file":
        ret = {"route": "nofile"}
        return json.dumps(ret)
    with open(name + '/'+ fileN, 'r', newline='') as csvFile:
        csv_file = csv.reader(csvFile)
        for content in csv_file:
            content=list(map(float,content))
            if len(content)!=0:
                data.append(content)
                traffic_feature.append(content[0:-2])
                traffic_target.append(content[-1])
    min_max_scaler = preprocessing.MinMaxScaler()
    traffic_feature = min_max_scaler.fit_transform(traffic_feature)
    # print('data=',data)
    # print('traffic_feature=',traffic_feature)
    # print('traffic_target=',traffic_target)
    dim = int(TSNE_dim["TSNE"])
    sklearn_TSNE = sklearnTSNE(n_components=dim)
    sklearn_transf = sklearn_TSNE.fit_transform(traffic_feature)
  
    with open(name + '/' + 'TSNE_'+fileN, 'w', newline='') as new_file:
         csv_writer = csv.writer(new_file)
         for line in sklearn_transf:
             # print(line)
             csv_writer.writerow(line)     
    ret = {"route": 'TSNE_'+fileN}
    return json.dumps(ret)

#kmeans聚类
@app.route('/kmeans/<name>', methods=['post', 'get'])
def kmeans(name):
    cluster_dim = request.get_json() #bytes
    data=[]
    cluster = int(cluster_dim["kmeans"])
    # print(name)
    fileN = db.searchFile(name)
    # print(fileN)
    fileN = fileN[0]['filename']
    # print(fileN)
    if fileN == "dont_have_file":
        ret = {"route": "nofile"}
        return json.dumps(ret)
    with open(name + '/'+ fileN, 'r', newline='') as csvFile:
        csv_file = csv.reader(csvFile)
        for content in csv_file:
            content=list(map(float,content))
            if len(content)!=0:
                data.append(content)            
    estimator = KMeans(n_clusters=cluster)#构造聚类器
    estimator.fit(data)#聚类
    label_pred = estimator.labels_ #获取聚类标签
    centroids = estimator.cluster_centers_ #获取聚类中心
    with open(name + '/' + 'kmeans_'+fileN, 'w', newline='') as new_file:
        csv_writer = csv.writer(new_file)
        csv_writer.writerow(["label"])
        for line in label_pred:
            # print(line)
            temp = str(line)
            # print(temp)
            csv_writer.writerow([temp])
        csv_writer.writerow(["clusterCenter"])
        for line in centroids:
            # print(line)
            csv_writer.writerow(line)

    ret = {"route": 'kmeans_'+fileN}
    return json.dumps(ret)

#单样本T检验
@app.route('/Ttest/<name>', methods=['post', 'get'])
def Ttest(name):
    cluster_dim = request.get_json() #bytes
    data=[]
    Tnum = float(cluster_dim["Tnum"])
    # print(name)
    fileN = db.searchFile(name)
    # print(fileN)
    fileN = fileN[0]['filename']
    # print(fileN)
    if fileN == "dont_have_file":
        ret = {"route": "nofile"}
        return json.dumps(ret)
    with open(name + '/'+ fileN, 'r', newline='') as csvFile:
        csv_file = csv.reader(csvFile)
        for content in csv_file:
            content=list(map(float,content))
            if len(content)!=0:
                data.append(content)
    result = stats.ttest_1samp(data, Tnum)
    print("result\n")
    print(result)
    with open(name + '/' + 'Ttest_'+fileN, 'w', newline='') as new_file:
        csv_writer = csv.writer(new_file)
        csv_writer.writerow(result[1])
        if result[1] < 0.05:
            csv_writer.writerow(["The difference is significant"])
        else:
            csv_writer.writerow(["The difference is not significant"])

    ret = {"route": 'Ttest_'+fileN}
    return json.dumps(ret)

#时序DTW分析
@app.route('/dtw/<name>', methods=['post', 'get'])
def dtw(name):
    fileN = db.searchFile(name)
    isfile = fileN[0]['filename']
    # print(fileN)
    if isfile == "dont_have_file":
        ret = {"route": "nofile"}
        return json.dumps(ret)
    if len(fileN) < 2:
        ret = {"route": '需要上传两个文件以进行时序列dtw分析'}
        return json.dumps(ret)
    data1=[]
    data2=[]
    file1 = fileN[0]['filename']
    file2 = fileN[1]['filename']
    csvFile1 = open(name + '/'+ file1)
    csv_file1 = csv.reader(csvFile1)
    for content in csv_file1:
        content = list(map(float,content))
        if len(content)!=0:
            data1.append(float(content[0]))
    csvFile1.close()
    print('data1=',data1)
    csvFile2 = open(name + '/'+ file2)
    csv_file2 = csv.reader(csvFile2)
    for content in csv_file2:
        content = list(map(float,content))
        if len(content)!=0:
            data2.append(float(content[0]))
    csvFile2.close()
    print('data2=',data2)
    r, c = len(data1), len(data2)
    D0 = zeros((r+1,c+1))
    D0[0,1:] = inf
    D0[1:,0] = inf
    D1 = D0[1:,1:]
    #浅复制
    # print D1
    for i in range(r):
        for j in range(c):
            D1[i,j] = euclidean_distances(data1[i],data2[j])
    #生成原始距离矩阵
    M = D1.copy()
    for i in range(r):
        for j in range(c):
            D1[i,j] += min(D0[i,j],D0[i,j+1],D0[i+1,j])
    #代码核心，动态计算最短距离
    i,j = array(D0.shape) - 2
    #最短路径
    # print i,j
    p,q = [i],[j]
    while(i>0 or j>0):
        tb = argmin((D0[i,j],D0[i,j+1],D0[i+1,j]))
        if tb==0 :
            i-=1
            j-=1
        elif tb==1 :
            i-=1
        else:
            j-=1
        p.insert(0,i)
        q.insert(0,j)
    print(M)
    #原始距离矩阵
    print(zip(p,q))
    #匹配路径过程
    print(D1)
    #Cost Matrix或者叫累积距离矩阵
    print(D1[-1,-1])
    dis = []
    dis.append(D1[-1,-1])
    print(dis)
    #序列距离
    with open(name + '/'+ 'DTW_result.csv', 'w', newline='') as new_file:
        csv_writer = csv.writer(new_file)
        csv_writer.writerow(["Cost Matrix"])
        for line in D1:
            # print(line)
            csv_writer.writerow(line)
        csv_writer.writerow(["Sequence distance"])
        csv_writer.writerow(dis)
    ret = {"route": 'DTW_result.csv'}
    return json.dumps(ret)

# 获取训练集数据
def get_train_dataset(name, filename):
    data = []
    with open(name + '/'+ filename, 'r', newline='') as csvFile:
        csv_file = csv.reader(csvFile)
        for content in csv_file:
            content = list(map(float,content))
            if len(content)!=0:
                data.append(content)
    random.shuffle(data)
    return data

# 获取测试集数据
def get_test_dataset(name, filename):
    data = []
    with open(name + '/'+ filename, 'r', newline='') as csvFile:
        csv_file = csv.reader(csvFile)
        for content in csv_file:
            content = list(map(float,content))
            if len(content)!=0:
                data.append(content)
    return data

        # content = list(map(float,content))
        # if len(content)!=0:
        #     data1.append(float(content[0]))

def get_nb_train_dataset(name, filename):
    data = []
    with open(name + '/'+ filename, 'r', newline='') as csvFile:
        csv_file = csv.reader(csvFile)
        for content in csv_file:
            if len(content)!=0:
                data.append(content)
    random.shuffle(data)
    return data

# 获取朴素贝叶斯处理文本分类问题的测试集
def get_nb_test_dataset(name, filename):
    data = []
    with open(name + '/'+ filename, 'r', newline='') as csvFile:
        csv_file = csv.reader(csvFile)
        for content in csv_file:
            if len(content)!=0:
                data.append(content[0])
    return data

# 将朴素贝叶斯处理文本分类问题的训练集划分为训练集和验证集
def nb_train_and_valid_data(data_):
    filesize = int(0.9 * len(data_))
    # 训练集和测试集的比例为9:1
    train_data_ = [each[0] for each in data_[:filesize]]
    train_target_ = [each[1] for each in data_[:filesize]]

    test_data_ = [each[0] for each in data_[filesize:]]
    test_target_ = [each[1] for each in data_[filesize:]]

    return train_data_, train_target_, test_data_, test_target_

def train_valid_data(data_):
    filesize = int(0.9 * len(data_))
    # 训练集和测试集的比例为9:1
    train_data_ = [each[0:-1] for each in data_[:filesize]]
    train_target_ = [each[-1] for each in data_[:filesize]]

    test_data_ = [each[0:-1] for each in data_[filesize:]]
    test_target_ = [each[-1] for each in data_[filesize:]]

    return train_data_, train_target_, test_data_, test_target_

#朴素贝叶斯处理文本分类问题
@app.route('/nb/<name>', methods=['post', 'get'])
def nb(name):
    print("first")
    fileN = db.searchFile(name)
    print("second")
    print(fileN)
    isfile = fileN[0]['filename']
    # print(fileN)
    if isfile == "dont_have_file":
        ret = {"route": "nofile"}
        return json.dumps(ret)
    if len(fileN) < 2:
        ret = {"route": '需要上传训练集和测试集以进行模型训练与测试'}
        return json.dumps(ret)
    train=[]
    test=[]
    train_file = fileN[0]['filename']
    test_file = fileN[1]['filename']
    print(train_file)
    print(test_file)
    train = get_nb_train_dataset(name, train_file)
    test = get_nb_test_dataset(name, test_file)
    print('train=',train)
    print('test=',test)
    train_data, train_target, valid_data, valid_target = nb_train_and_valid_data(train)
    nbc_6 = Pipeline([
        ('vect', TfidfVectorizer(
                             
        )),
        ('clf', MultinomialNB(alpha=1.0)),
    ])
    nbc_6.fit(train_data, train_target)    #训练我们的多项式模型贝叶斯分类器
    valid = nbc_6.predict(valid_data)  #在测试集上预测结果
    print(valid)
    count = 0                                      #统计预测正确的结果个数
    for left , right in zip(valid, valid_target):
        if left == right:
            count += 1
    print(count/len(valid_target))
    predict = nbc_6.predict(test)
    print(predict)
    p_list = []
    for p in predict:
        list_tem = []
        list_tem.append(p)
        p_list.append(list_tem)
    print(p_list)
    with open(name + '/'+ 'NaiveBayes.csv', 'w', newline='') as new_file:
        csv_writer = csv.writer(new_file)
        for line in p_list:
            # print(line)
            csv_writer.writerow(list(line))
    ret = {"route": 'NaiveBayes.csv'}
    return json.dumps(ret)

#决策树处理分类问题
@app.route('/decisionTree/<name>', methods=['post', 'get'])
def decisionTree(name):
    fileN = db.searchFile(name)
    isfile = fileN[0]['filename']
    # print(fileN)
    if isfile == "dont_have_file":
        ret = {"route": "nofile"}
        return json.dumps(ret)
    if len(fileN) < 2:
        ret = {"route": '需要上传训练集和测试集以进行模型训练与测试'}
        return json.dumps(ret)
    train=[]
    test=[]
    train_file = fileN[0]['filename']
    test_file = fileN[1]['filename']
    print(train_file)
    print(test_file)
    train = get_train_dataset(name, train_file)
    test = get_test_dataset(name, test_file)
    # print('test=',test)
    train_data, train_target, valid_data, valid_target = train_valid_data(train)
    model = tree.DecisionTreeClassifier(criterion='gini')
    # print('train_data=',train_data)
    # print('train_target=',train_target)
    model.fit(train_data,train_target)
    print("train score:", model.score(train_data, train_target))
    print("valid score:", model.score(valid_data, valid_target))
    predict = model.predict(test)
    print(predict)
    p_list = []
    for p in predict:
        list_tem = []
        list_tem.append(p)
        p_list.append(list_tem)
    print(p_list)
    with open(name + '/'+ 'DecisionTree.csv', 'w', newline='') as new_file:
        csv_writer = csv.writer(new_file)
        for line in p_list:
            # print(line)
            csv_writer.writerow(list(line))
    ret = {"route": 'DecisionTree.csv'}
    return json.dumps(ret)

#随机森林处理分类问题
@app.route('/randomForest/<name>', methods=['post', 'get'])
def randomForest(name):
    fileN = db.searchFile(name)
    isfile = fileN[0]['filename']
    # print(fileN)
    if isfile == "dont_have_file":
        ret = {"route": "nofile"}
        return json.dumps(ret)
    if len(fileN) < 2:
        ret = {"route": '需要上传训练集和测试集以进行模型训练与测试'}
        return json.dumps(ret)
    train=[]
    test=[]
    train_file = fileN[0]['filename']
    test_file = fileN[1]['filename']
    print(train_file)
    print(test_file)
    train = get_train_dataset(name, train_file)
    test = get_test_dataset(name, test_file)
    print('train=',train)
    print('test=',test)
    train_data, train_target, valid_data, valid_target = train_valid_data(train)
    # Create the model with 100 trees
    model = RandomForestClassifier(n_estimators=100, 
                               bootstrap = True,
                               max_features = 'sqrt')
    # Fit on training data
    model.fit(train_data,train_target)
    print("train score:", model.score(train_data, train_target))
    print("valid score:", model.score(valid_data, valid_target))
    predict = model.predict(test)
    print(predict)
    p_list = []
    for p in predict:
        list_tem = []
        list_tem.append(p)
        p_list.append(list_tem)
    print(p_list)
    with open(name + '/'+ 'RandomForest.csv', 'w', newline='') as new_file:
        csv_writer = csv.writer(new_file)
        for line in p_list:
            # print(line)
            csv_writer.writerow(list(line))
    ret = {"route": 'RandomForest.csv'}
    return json.dumps(ret)

#svm 支持向量机处理分类问题
@app.route('/svm/<name>', methods=['post', 'get'])
def svm(name):
    fileN = db.searchFile(name)
    isfile = fileN[0]['filename']
    # print(fileN)
    if isfile == "dont_have_file":
        ret = {"route": "nofile"}
        return json.dumps(ret)
    if len(fileN) < 2:
        ret = {"route": '需要上传训练集和测试集以进行模型训练与测试'}
        return json.dumps(ret)
    train=[]
    test=[]
    train_file = fileN[0]['filename']
    test_file = fileN[1]['filename']
    train = get_train_dataset(name, train_file)
    test = get_test_dataset(name, test_file)
    train_data, train_target, valid_data, valid_target = train_valid_data(train)
    svclassifier = SVC(kernel='poly', degree=8)
    svclassifier.fit(train_data, train_target)
    print("train score:", svclassifier.score(train_data, train_target))
    print("valid score:", svclassifier.score(valid_data, valid_target))
    predict = svclassifier.predict(test)
    print(predict)
    p_list = []
    for p in predict:
        list_tem = []
        list_tem.append(p)
        p_list.append(list_tem)
    print(p_list)
    with open(name + '/'+ 'svm.csv', 'w', newline='') as new_file:
        csv_writer = csv.writer(new_file)
        for line in p_list:
            # print(line)
            csv_writer.writerow(list(line))
    ret = {"route": 'svm.csv'}
    return json.dumps(ret)

#svr 支持向量回归机处理回归问题
@app.route('/svr/<name>', methods=['post', 'get'])
def svr(name):
    fileN = db.searchFile(name)
    isfile = fileN[0]['filename']
    # print(fileN)
    if isfile == "dont_have_file":
        ret = {"route": "nofile"}
        return json.dumps(ret)
    if len(fileN) < 2:
        ret = {"route": '需要上传训练集和测试集以进行模型训练与测试'}
        return json.dumps(ret)
    train=[]
    test=[]
    train_file = fileN[0]['filename']
    test_file = fileN[1]['filename']
    print(train_file)
    print(test_file)
    train = get_train_dataset(name, train_file)
    test = get_test_dataset(name, test_file)
    print('train=',train)
    print('test=',test)
    train_data, train_target, valid_data, valid_target = train_valid_data(train)
    svrlassifier = SVR(kernel = 'rbf', c = 20)
    # Fit on training data
    svrlassifier.fit(train_data,train_target)
    print("train score:", svrlassifier.score(train_data, train_target))
    print("valid score:", svrlassifier.score(valid_data, valid_target))
    predict = svrlassifier.predict(test)
    print(predict)
    p_list = []
    for p in predict:
        list_tem = []
        list_tem.append(p)
        p_list.append(list_tem)
    print(p_list)
    with open(name + '/'+ 'svr.csv', 'w', newline='') as new_file:
        csv_writer = csv.writer(new_file)
        for line in p_list:
            # print(line)
            csv_writer.writerow(list(line))
    ret = {"route": 'svr.csv'}
    return json.dumps(ret)

#linearRegression 多元线性回归
@app.route('/linearRegression/<name>', methods=['post', 'get'])
def linearRegression(name):
    fileN = db.searchFile(name)
    isfile = fileN[0]['filename']
    # print(fileN)
    if isfile == "dont_have_file":
        ret = {"route": "nofile"}
        return json.dumps(ret)
    if len(fileN) < 2:
        ret = {"route": '需要上传训练集和测试集以进行模型训练与测试'}
        return json.dumps(ret)
    train=[]
    test=[]
    train_file = fileN[0]['filename']
    test_file = fileN[1]['filename']
    print(train_file)
    print(test_file)
    train = get_train_dataset(name, train_file)
    test = get_test_dataset(name, test_file)
    print('train=',train)
    print("\n")
    train_data, train_target, valid_data, valid_target = train_valid_data(train)
    linreg = LinearRegression()
    model = linreg.fit(train_data, train_target)
    print (model)
    # 训练后模型截距
    print (linreg.intercept_)
    # 训练后模型权重（特征个数无变化）
    print (linreg.coef_)
    predict = linreg.predict(test)
    p_list = []
    for p in predict:
        list_tem = []
        list_tem.append(p)
        p_list.append(list_tem)
    print(p_list)
    with open(name + '/'+ 'linearRegression.csv', 'w', newline='') as new_file:
        csv_writer = csv.writer(new_file)
        for line in p_list:
            # print(line)
            csv_writer.writerow(list(line))
    ret = {"route": 'linearRegression.csv'}
    return json.dumps(ret)

# LogisticRegression 逻辑回归处理分类问题
@app.route('/logisticRegression/<name>', methods=['post', 'get'])
def logisticRegression(name):
    fileN = db.searchFile(name)
    isfile = fileN[0]['filename']
    # print(fileN)
    if isfile == "dont_have_file":
        ret = {"route": "nofile"}
        return json.dumps(ret)
    if len(fileN) < 2:
        ret = {"route": '需要上传训练集和测试集以进行模型训练与测试'}
        return json.dumps(ret)
    train=[]
    test=[]
    train_file = fileN[0]['filename']
    test_file = fileN[1]['filename']
    print(train_file)
    print(test_file)
    train = get_train_dataset(name, train_file)
    test = get_test_dataset(name, test_file)
    # print('train=',train)
    # print('test=',test)
    train_data, train_target, valid_data, valid_target = train_valid_data(train)
    logreg = LogisticRegression()
    model = logreg.fit(train_data, train_target)
    print (model)
    predict = model.predict(test)
    print("train score:", model.score(train_data, train_target))
    print("valid score:", model.score(valid_data, valid_target))
    p_list = []
    for p in predict:
        list_tem = []
        list_tem.append(p)
        p_list.append(list_tem)
    with open(name + '/'+ 'logisticRegression.csv', 'w', newline='') as new_file:
        csv_writer = csv.writer(new_file)
        for line in p_list:
            # print(line)
            csv_writer.writerow(list(line))
    ret = {"route": 'logisticRegression.csv'}
    return json.dumps(ret)

#wordcloud词云生成
@app.route('/wordCloud/<name>', methods=['post', 'get'])
def workCloud(name):
    # print(name)
    fileN = db.searchFile(name)
    # print(fileN)
    fileN = fileN[0]['filename']
    # print(fileN)
    if fileN == "dont_have_file":
        ret = {"route": "nofile"}
        return json.dumps(ret)
    txtfile = fileN
    print(txtfile)
    f = open(name + '/'+ txtfile,encoding = 'gb18030')
    text = f.read()
    print(text)
    f.close()
    # Read the whole text.
    # #获取当前的项目文件加的路径
    # d=path.dirname(__file__) 
    # text = open(path.join(d, txtfile))
    # 结巴分词
    wordlist = jieba.cut(text, cut_all=True)
    wl = " ".join(wordlist)
    print(wl)#输出分词之后的txt
    coloring = np.array(Image.open("wordcloud/background.jpg"))

    # 设置停用词
    stopwords = set(STOPWORDS)
    stopwords.add("said")

    # 你可以通过 mask 参数 来设置词云形状
    # wc = WordCloud(background_color="white", max_words=2000, mask=coloring,
    #                 max_font_size=50, random_state=42,font_path='fangsong_GB2312.ttf')
    wc = WordCloud(background_color="white", max_words=2000, mask=coloring,
                max_font_size=50, random_state=42,font_path='Hiragino Sans GB.ttc')
    wc.generate(wl)

    # create coloring from image
    image_colors = ImageColorGenerator(coloring)
    # show
    # 在只设置mask的情况下,你将会得到一个拥有图片形状的词云
    # plt.imshow(wc, interpolation="bilinear")
    # plt.axis("off")
    # plt.figure()
    # plt.show()
    wc.to_file(name + '/'+ "workcloud.png")
    ret = {"route": 'workcloud.png'}
    return json.dumps(ret)

#pdf2word 文献pdf内容提取
@app.route('/pdf2word/<name>', methods=['post', 'get'])
def pdf2word(name):
    # pip install pdfminer3k
    # pip install python_docx
    fileN = db.searchFile(name)
    # print(fileN)
    fileN = fileN[0]['filename']
    # print(fileN)
    if fileN == "dont_have_file":
        ret = {"route": "nofile"}
        return json.dumps(ret)
    pdfname = fileN
    document = Document()
    pdf=os.open(name + '/'+ pdfname,os.O_RDWR )
    fn = open(pdf,'rb')
    parser = PDFParser(fn)
    doc = PDFDocument()
    parser.set_document(doc)
    doc.set_parser(parser)
    resource = PDFResourceManager()
    laparams = LAParams()
    device = PDFPageAggregator(resource,laparams=laparams)
    interpreter = PDFPageInterpreter(resource,device)
    for i in doc.get_pages():
        interpreter.process_page(i)
        layout = device.get_result()
        for out in layout:
            if hasattr(out,"get_text"):
                content = out.get_text().replace(u'\n', u'') 
                document.add_paragraph(
                    content, style='ListBullet'   
                )
            document.save(name + '/'+ pdfname[0:-4]+'.docx')
    print ('处理完成')
    ret = {"route": pdfname[0:-4]+'.docx'}
    return json.dumps(ret)

#自定义列向量插值函数
def ploy(s,n,k=3):
    y=s[list(range(n-k,n))+list(range(n+1,n+1+k))]#取数
    y=y[y.notnull()]
    return lagrange(y.index,list(y))(n)

#poly 拉格朗日插值处理文本数据缺失值
@app.route('/polydata/<name>', methods=['post', 'get'])
def polydata(name):
    fileN = db.searchFile(name)
    # print(fileN)
    fileN = fileN[0]['filename']
    # print(fileN)
    if fileN == "dont_have_file":
        ret = {"route": "nofile"}
        return json.dumps(ret)
    csv_file = fileN
    csv_data = pd.read_csv(name + '/'+ csv_file, low_memory = False)#防止弹出警告
    data = pd.DataFrame(csv_data)

    for i in data.columns:
        for j in range(len(data)):
            if(np.isnan(data[i][j])):
                data[i][j]=ploy(data[i],j)
                # print(data[i][j])

    data.to_csv(name + '/'+ csv_file[0:-4]+'.csv', index=False, header=False )
    print ('处理完成')
    ret = {"route": csv_file[0:-4]+'.csv'}
    return json.dumps(ret)

#计算两个序列的最大信息系数MIC
@app.route('/MIC/<name>', methods=['post', 'get'])
def MIC(name):
    # pip install minepy
    fileN = db.searchFile(name)
    isfile = fileN[0]['filename']
    # print(fileN)
    if isfile == "dont_have_file":
        ret = {"route": "nofile"}
        return json.dumps(ret)

    if len(fileN) < 2:
        ret = {"route": '需要上传两个文件以进行MIC计算'}
        return json.dumps(ret)
    x=[]
    y=[]
    file1 = fileN[0]['filename']
    file2 = fileN[1]['filename']
    csvFile1 = open(name + '/'+ file1, encoding='utf-8-sig')
    csv_file1 = csv.reader(csvFile1)
    for content in csv_file1:
        print(content)
        content = list(map(float,content))
        if len(content)!=0:
            x.append(float(content[0]))
    csvFile1.close()
    print('x=',x)
    csvFile2 = open(name + '/'+ file2, encoding='utf-8-sig')
    csv_file2 = csv.reader(csvFile2)
    for content in csv_file2:
        content = list(map(float,content))
        if len(content)!=0:
            y.append(float(content[0]))
    csvFile2.close()
    print('y=',y)
    mine = MINE(alpha = 0.6, c = 15)
    mine.compute_score(x, y)
    print("MIC", mine.mic())
    #将MIC值写入文件
    with open(name + '/'+ 'MIC_result.csv', 'w', newline='') as new_file:
        csv_writer = csv.writer(new_file)
        csv_writer.writerow(["MIC result"])
        data = []
        data.append(str(mine.mic()))
        csv_writer.writerow(data)
    ret = {"route": 'MIC_result.csv'}
    return json.dumps(ret)

#下载处理完成的文件
@app.route('/download/<name>/<filename>',  methods=['post', 'get'])
def download(name,filename):
    # data = request.get_json()
    # file = data['file']
    # print("hhh\n")
    # print(file)
    # 需要知道2个参数, 第1个参数是本地目录的path, 第2个参数是文件名(带扩展名)
    # directory = os.getcwd()  # 假设在当前目录
    # response = make_response(send_from_directory(directory, file, as_attachment=True))
    # response.headers["Content-Disposition"] = "attachment; file={}".format(file_name.encode().decode('latin-1'))
    # return response
    print(name)
    print(filename)
    directory = os.getcwd()  # 假设在当前目录
    print(directory)
    path = name + '/' + filename
    print(path)
    return send_from_directory(directory, path, as_attachment=True)
    # ret = {"status": "success"}
    # return json.dumps(ret)

#移除用户上传的文件
@app.route('/removeFile/<name>', methods=["post"])
def removeFile(name):
    file = request.get_json()
    fileName = file['name']
    if os.path.exists(name + '/' + fileName): # 如果文件存在
        db.deleteFile(name, fileName)
        os.remove(name + '/' + fileName) # 则删除
        ret = {"status": "success"}
    else:
        print('no such file:%s'%fileName)
        ret = {"status": "no such file"}
    return json.dumps(ret)

#获取问卷列表信息
@app.route('/list/', methods=['post', 'get'])
def showList():
    user = request.get_json()
    # user = user["user"]
    print(user)
    result = db.get_list_info(user)
    ret = []
    for i in result:
        # if i['randomNum'] != '0':
        #     q_final = []
        #     for q in i['question']:
        #         if q['isRandom'] == 0:
        #             q_final.append(q)
        #     for cnt in range(1, int(i['randomNum'])+1):
        #         print("cnt= ")
        #         print(cnt)
        #         q_list = []
        #         for q in i['question']:
        #             print(q['isRandom'])
        #             print("hhh")
        #             if int(q['isRandom']) == cnt:
        #                 print("shabi")
        #                 q_list.append(q)
        #                 print(q_list)
        #         print(len(q_list))
        #         rand = random.randint(0,len(q_list)-1)
        #         q_final.append(q_list[rand])
        #         print("q_final\n")
        #         print(q_final)
        #     i['question'] = q_final
        # else:
        #     pass
        ret.append({
            "user": i['user'],
            "num": i['num'],
            "title": i['title'],
            "time": i['time'],
            "state": i['state'],
            "stateTitle": i['stateTitle'],
            "checked": i['checked'],
            "question": i['question']
        })
    return json.dumps(ret)

#获取问卷列表信息
@app.route('/fillList/', methods=['post', 'get'])
def fillList():
    print("come")
    user = request.get_json()
    result = db.fill_list(user['user'])
    ret = []
    for i in result:
        if i['num'] == int(user["num"]):
            ret.append({
                "user": i['user'],
                "num": i['num'],
                "title": i['title'],
                "time": i['time'],
                "state": i['state'],
                "stateTitle": i['stateTitle'],
                "checked": i['checked'],
                "question": i['question']
            })
    print("ret")
    print(ret)
    return json.dumps(ret)

#新增问卷
@app.route('/editList/',methods=['post', 'get'])
def addList():
    data = request.get_json() #bytes
    print("data_receive\n")
    print(data)    
    db.addList(data)
    ret = {"status": "success"}
    return json.dumps(ret)

#删除问卷
@app.route('/deleteList/',methods=['post', 'get'])
def deleteList():
    # print("hhhhhhh\n")
    # temp = request.get_data()
    # num = bytes.decode(temp)
    # num = json.loads(num)
    # print(type(num))
    data = request.get_json() #bytes
    db.deleteList(data)
    ret = {"status": "success"}
    return json.dumps(ret)

#填写问卷
@app.route('/addChoose/',methods=['post', 'get'])
def addChoose():
    data = request.get_json() #bytes
    # print("data_receive\n")
    # print(data)    
    db.addChoose(data)
    ret = {"status": "success"}
    return json.dumps(ret)

#获取问卷回答信息
@app.route('/showData/', methods=['post', 'get'])
def showData():
    data = request.get_json() #bytes
    print(data)
    choose_result = db.get_choose_info(data['user'])
    choose = []
    for i in choose_result:
        if i['num'] == int(data["num"]):
            choose.append({
                "num": i['num'],
                "question": i['question']
            })
    print("choose")
    print(choose)
    result = db.fill_list(data['user'])
    ques = []
    for i in result:
        if i['num'] == int(data["num"]):
            ques.append({
                "num": i['num'],
                "question": i['question']
            })
    ques = ques[0]
    print("ques")
    print(ques)
    res = []
    for q in ques['question']:
        if q['type'] == 'radio' or q['type'] == 'checkbox' or q['type'] == 'rate':
            # q = 'num': 'Q1', 'title': '单选题', 'type': 'radio', 'isNeed': True, 'options': ['1', '2', '3'], 'isHidden': False
            q_dict = dict()
            chos = q['options']
            for cho in chos:
                q_dict[cho] = 0
            count_dict = {
                q['num'] : q_dict
            }
            q_num = q['num']
            for c in choose:
                if(type(c['question'][q_num]) == str):
                    ans = c['question'][q_num]
                    count_dict[q_num][ans] = count_dict[q_num][ans] + 1
                else:
                    for ans in c['question'][q_num]:
                        count_dict[q_num][ans] = count_dict[q_num][ans] + 1
            # print("count_dict = \n")
            # print(count_dict)
            res.append(count_dict)
        elif q['type'] == 'textarea':
            # 'num': 'Q3', 'title': '文本题', 'type': 'textarea', 'isNeed': True
            q_list = []
            count_dict = {
                q['num'] : q_list
            }
            for c in choose:
                q_list.append(c['question'][q['num']])
            res.append(count_dict)
        else:
            pass
    print("res")
    print(res)
    return json.dumps(res)

@app.route('/keywordExtraction/', methods=["POST"])
def keywordExtraction():
    data = request.get_json() #bytes
    # print("data_receive\n")
    # print(data)
    textarea = data["data"];
    topK = 6;
    jieba.analyse.set_stop_words("data/stopWord.txt") # 加载自定义停用词表
    keywords = jieba.analyse.textrank(textarea, topK=topK, allowPOS=('n','nz','v','vd','vn','l','a','d'))
    word_split = "  ".join(keywords)
    print(word_split)
    text_key = {}
    text_key[textarea] = word_split
    db.addkeywordExtraction(text_key)
    return json.dumps(word_split)

@app.route('/getWordFrequency/', methods=["POST"])
def getWordFrequency():
    text = request.get_json() #bytes
    # print("data_receive\n")
    # print(data)
    data = text["data"];
    print(data)
    #Get Hasgtag WordBox
    data = data.lower()
    data = re.sub(r'\|\~|\`|\!|\$|\%|\^|\&|\*|\(|\)|\-|\_|\+|\=|\||\\|\[|\]|\{|\}|\;|\:|\"|\'|\,|\<|\.|\>|\/|\?', " ", data)
    print(data)
    wordsBox = data.strip().split()
    print(wordsBox)
    wordsBoxofWord = []
    for word in wordsBox:
        #get word
        if re.match(r'([a-zA-Z0-9])', word):
            wordsBoxofWord.append(word)
    #Get frequency
    c = collections.Counter(wordsBoxofWord)
    # return json.dumps(word_split)
    ret = {"status": "success"}
    return json.dumps(c)

@app.route('/freAnlysisPuc/', methods=["POST"])
def freAnlysisPuc():
    text = request.get_json() #bytes
    data = text["data"];
    print(data)
    txt1 = data
    txt1 = txt1.replace('\n', '')  # 删掉换行符
    txt1 = txt1.lower()
    for ch in 'abcdefghijklmnopqrstuvwxyz1234567890 ':
        txt1 = txt1.replace(ch, '')
    mylist = list(txt1)
    mycount = collections.Counter(mylist)
    results = []
    for key, val in mycount.most_common():
        results.append((key, val))
    sorted(results, key=lambda x: x[1], reverse=True)
    print(results)
    return json.dumps(results)

@app.route('/getSentiment/', methods=["POST"])
def getSentiment():
    text = request.get_json() #bytes
    data = text["data"];
    print(data)
    sentiment = TextBlob(data)
    polarityResult = sentiment.polarity
    subjectivityResult = sentiment.subjectivity
    print(polarityResult)
    print(subjectivityResult)
    results = {}
    new_polarityResult = round( polarityResult , 3 )
    new_subjectivityResult = round( subjectivityResult , 3 )
    results['polarity'] = new_polarityResult
    results['subjectivity'] = new_subjectivityResult
    return json.dumps(results)


# @app.route('/searchArticle/', methods=["POST"])
# def searchArticle():
#     data = request.get_json() #bytes
#     article = data["data"];
#     print(article)
#     ws_api = wechatsogou.WechatSogouAPI(captcha_break_time=3)
#     data_receive = ws_api.search_article(article, identify_image_callback=identify_image_callback_ruokuai_sogou)
#     results = []
#     for i in data_receive:
#         print(i)
#         results.append(i)
#     return json.dumps(results)

#登陆界面，检查用户名是否存在，若不存在则直接注册
@app.route('/login/',methods=["POST"])
def login():
    data = request.get_json() #bytes
    print("data_receive")
    print(data)    
    state = db.login(data)
    print("state")
    print(state["state"])
    if state["state"] == "register_success":
        os.makedirs(data["name"])
    ret = {"status": state["state"]}
    return json.dumps(ret)

@app.route('/get_form_completion/<user>/<num>/<qid>', methods=['GET'])
def get_form_completion(user,num,qid):
    result = db.user_result(user,num)
    ret = []
    for i in result:
        ret.append(
            i['question'][qid]
        )
    print(ret)
    return json.dumps(ret)

@app.route('/get_mic_data', methods=['GET','POST'])
def MICvalue():
    choice_user = request.get_json()  # 获取前端用户选择的数据
    flag = True
    # data = request.get_json() #bytes
    # print(data)
    choice0 = {}
    choice1 = {}
    # choice[0]['db'] = data[0][db]
    # choice[0]['col'] = data[0][col]
    # choice[0]['field'] = data[0][field]
    # choice[1]['db'] = data[1][db]
    # choice[1]['col'] = data[1][col]
    # choice[1]['field'] = data[1][field]
    choice0['db'] = choice_user[0][0]
    choice0['col'] = choice_user[0][1]
    choice0['field'] = choice_user[0][2]
    choice1['db'] = choice_user[1][0]
    choice1['col'] = choice_user[1][1]
    choice1['field'] = choice_user[1][2]
    print("choice0", choice0)
    print("choice1", choice1)
    # choice0['db'] = 'EpidemicData'
    # choice0['col'] = '上海'
    # choice0['field'] = '新增确诊'
    # choice1['db'] = 'EpidemicData'
    # choice1['col'] = '河北'
    # choice1['field'] = '新增确诊'
    # print(choice0)
    # print(choice1)

    # 获取数据
    # client = MongoClient("10.72.100.5",8027,username='double',password='double')
    client = MongoClient("10.72.100.5",8027)
    db = client.admin
    db.authenticate("double", "double")
    conn = MongoClient(host='mongodb://10.72.100.5:8027/'+'admin',username='double',password='double')
    database = conn[choice0['db']]
    collection0 = database[choice0['col']]
    results0 = collection0.find({},{choice0['field']:1,"_id":0}).sort("_id",pymongo.ASCENDING)   # 按照_id排序
    collection1 = database[choice1['col']]
    results1 = collection1.find({},{choice1['field']:1,"_id":0}).sort("_id",pymongo.ASCENDING)   # 按照_id排序
    # 1表示显示此字段，0表示不显示此字段，默认会显示_id
    rawdata0 = []
    rawdata1 = []
    for result in results0:
        rawdata0.append(result[choice0['field']])
    for result in results1:
        rawdata1.append(result[choice1['field']])

    # 清理数据
    for i in range(len(rawdata0)-1,-1,-1):   # 假定rawdata0与rawdata1的长度相同
        if rawdata0[i] and rawdata1[i]:
            try:   # 将数字形式的数据转换为浮点数
                rawdata0[i] = float(rawdata0[i])
                rawdata1[i] = float(rawdata1[i])
            except ValueError:
                flag = False   # 存在非数值字段
        else:
            del rawdata0[i]
            del rawdata1[i]

    print("rawdata0", rawdata0)
    print("rawdata1", rawdata1)
    # 计算MIC
    m = MINE()
    if rawdata0:   # 当rawdata0与rawdata1不为空时
        if flag:
            # 将数据映射到[0,1]区间
            min_max_scaler = MinMaxScaler()
            data1_std = min_max_scaler.fit_transform(np.array(rawdata0).reshape(-1, 1))
            data2_std = min_max_scaler.fit_transform(np.array(rawdata1).reshape(-1, 1))
            data1 = data1_std.reshape(1,-1)[0]
            data2 = data2_std.reshape(1,-1)[0]
            m.compute_score(data1,data2)
            # str(m.mic())
            return json.dumps(m.mic())
        else:
            return "请选取数值字段"
    else:
        return "您所选取的两个字段无对应数据"

# 获取options
@app.route('/getOptions', methods=['post', 'get'])
def getOptions():
    client = MongoClient("10.72.100.5",8027)
    db = client.admin
    db.authenticate("double", "double")
    dblist = client.list_database_names()   # 服务器上除系统数据库外的所有数据库
    print('所有的数据库：', dblist)
    j = 0 
    for i in range(len(dblist)):
        if dblist[j] == 'admin':
            dblist.pop(j)
        elif dblist[j] == 'config':
            dblist.pop(j)
        elif dblist[j] == 'local':
            dblist.pop(j)
        else:
            j += 1
    print(dblist)
    options_list = []
    for i in range(len(dblist)):    
        db = dblist[i]
        database = client[db]
        collection_list = database.list_collection_names()   # 指定数据库中的所有集合
        
        child_list0 = []
        for j in range(len(collection_list)):
            coll = collection_list[j]
            collection = database[coll]
            document = collection.find_one()
            field_list = list(document.keys())[1:]   # 指定数据库中的所有字段(除了"_id")
            
            child_list1 = []
            for k in range(len(field_list)):
                child_list1.append({'value':field_list[k],'label':field_list[k]})
                
            child_list0.append({'value':collection_list[j],'label':collection_list[j],'children':child_list1})
            
        options_list.append({'value':dblist[i],'label':dblist[i],'children':child_list0})
   
    return json.dumps(options_list)

#博文信息
# @app.route('/API/getStu/<id>')
# def getStu(id=id):
#     result = db.get_stu_info(id)
#     #id,title,type,pubTime
#     ret = []
#     for i in result:
#         ret.append({
#             "id": i[0],
#             "name": i[1],
#             "age": i[2],
#             "sex": i[3],
#             "birthtime": i[4].strftime("%Y-%m-%d"),
#             "class": i[5],
#             "address": i[6],
#             "tel":i[7]
#         })
#     print "ret"+json.dumps(ret)
#     return json.dumps(ret)

# @app.route('/favicon.ico')
# def favicon(id=id):
#     return app.send_static_file("./static/favicon.ico")

# # token <=> password id
# # verifyToken return isUser userId
# #验证账户
# def verifyToken(token):
#     SQLresult = db.verifyToken()
#     if token == None:
#         return False, "null"
#     for i in SQLresult:
#         print i
#         if token == genCookie(i[0]):
#             return True, i[1]
#     return False, "null"


# #加密
# def genCookie(passMd5):
#     today = datetime.date.today()
#     Md5 = hashlib.md5()
#     Md5.update(passMd5 + today.strftime("%Y/%m/%d"))
#     Md5hex = Md5.hexdigest()
#     return Md5hex


if __name__ == '__main__':
    app.run(host="127.0.0.1", debug=False, port=8025)
